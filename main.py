#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import re
import time
from aip import AipNlp
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from pyhanlp import *
from ratelimiter import RateLimiter
import configparser
cfg = configparser.ConfigParser()

cfg.read('app.config')

APP_ID = cfg.get('baidu', 'APP_ID')
API_KEY = cfg.get('baidu', 'API_KEY')
SECRET_KEY = cfg.get('baidu', 'SECRET_KEY')

print(APP_ID, API_KEY, SECRET_KEY)
client = AipNlp(APP_ID, API_KEY, SECRET_KEY)


def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def table_print(block):
    table = block
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                text = paragraph.text
                segment = HanLP.newSegment('crf').enableNameRecognize(True)
                res = segment.seg(text)
                # res = client.lexer(text)
                print(res)


document = Document('test.docx')

paragraphs = [paragraph.text
              for paragraph in document.paragraphs]
tables = document.tables

for index, block in enumerate(iter_block_items(document)):
    if isinstance(block, Paragraph):
        print(block.text)
    elif isinstance(block, Table):
        table_print(block)
        last_block = list(iter_block_items(document))[index - 1]
        if isinstance(last_block, Paragraph):
            print(last_block.text)

# res = client.lexer('马云')
# print(res)

# rate_limiter = RateLimiter(max_calls=3, period=1)

# for i in range(100):
#     with rate_limiter:
#         res = client.lexer('马云')
#         print(res)

# def limited(until):
#     duration = int(round(until - time.time()))
#     print('Rate limited, sleeping for {:d} seconds'.format(duration))


# rate_limiter = RateLimiter(max_calls=2, period=1, callback=limited)


# for i in range(100):
#     with rate_limiter:
#         print('Iteration', i)
